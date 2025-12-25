"""
SSH Guardian v3.1 - Dashboard Content Routes
API endpoints for guide steps and thesis content
"""
import sys
import io
import re
from pathlib import Path
from flask import Blueprint, jsonify, send_file
from html import unescape

PROJECT_ROOT = Path(__file__).parent.parent.parent.parent
sys.path.append(str(PROJECT_ROOT / "dbs"))
sys.path.append(str(PROJECT_ROOT / "src" / "core"))

from connection import get_connection
from cache import get_cache, cache_key

# Import python-docx for DOCX export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

dashboard_content_routes = Blueprint('dashboard_content', __name__)

# Cache TTL - Extended for static content
GUIDE_CACHE_TTL = 86400  # 24 hours


# ==============================================================================
# GUIDE STEPS ENDPOINTS (v3.1 schema)
# ==============================================================================

@dashboard_content_routes.route('/guide/full', methods=['GET'])
def get_full_guide():
    """Get complete guide data including all steps - v3.1 schema"""
    try:
        cache = get_cache()
        cache_k = cache_key('guide', 'full', 'all')
        cached = cache.get(cache_k)
        if cached is not None:
            return jsonify({'success': True, 'data': cached, 'from_cache': True})

        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        # v3.1 schema: step_key, title, description, step_order, is_required, icon, action_url
        cursor.execute("""
            SELECT id, step_key, title, description, step_order, icon, action_url
            FROM guide_steps
            ORDER BY step_order ASC
        """)
        db_steps = cursor.fetchall()

        cursor.close()
        conn.close()

        # Map to frontend expected format with default content
        steps = []
        for db_step in db_steps:
            step = {
                'id': db_step['id'],
                'step_number': db_step['step_order'],
                'step_key': db_step['step_key'],
                'title': db_step['title'],
                'subtitle': db_step['description'],
                'icon': db_step['icon'] or '&#128214;',
                'content_html': _get_default_step_content(db_step['step_key']),
                'tips_html': _get_default_step_tips(db_step['step_key'])
            }
            steps.append(step)

        result = {
            'total_steps': len(steps),
            'steps': steps
        }

        cache.set(cache_k, result, GUIDE_CACHE_TTL)

        return jsonify({'success': True, 'data': result, 'from_cache': False})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@dashboard_content_routes.route('/guide/steps', methods=['GET'])
def get_guide_steps():
    """Get all guide steps - v3.1 schema"""
    try:
        cache = get_cache()
        cache_k = cache_key('guide', 'steps', 'all')
        cached = cache.get(cache_k)
        if cached is not None:
            return jsonify({'success': True, 'data': cached, 'from_cache': True})

        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("""
            SELECT id, step_key, title, description, step_order, icon, action_url
            FROM guide_steps
            ORDER BY step_order ASC
        """)
        db_steps = cursor.fetchall()

        cursor.close()
        conn.close()

        # Map to frontend expected format
        steps = []
        for db_step in db_steps:
            step = {
                'id': db_step['id'],
                'step_number': db_step['step_order'],
                'step_key': db_step['step_key'],
                'title': db_step['title'],
                'subtitle': db_step['description'],
                'icon': db_step['icon'] or '&#128214;',
                'content_html': _get_default_step_content(db_step['step_key']),
                'tips_html': _get_default_step_tips(db_step['step_key'])
            }
            steps.append(step)

        cache.set(cache_k, steps, GUIDE_CACHE_TTL)

        return jsonify({'success': True, 'data': steps, 'from_cache': False})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ==============================================================================
# THESIS ENDPOINTS (v3.1: restored with actual database tables)
# ==============================================================================

@dashboard_content_routes.route('/thesis/full', methods=['GET'])
def get_full_thesis():
    """Get complete thesis data from database"""
    try:
        cache = get_cache()
        cache_k = cache_key('thesis', 'full', 'all')
        cached = cache.get(cache_k)
        if cached is not None:
            return jsonify({'success': True, 'data': cached, 'from_cache': True})

        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        # Get metadata
        cursor.execute("SELECT meta_key, meta_value FROM thesis_metadata")
        metadata_rows = cursor.fetchall()
        metadata = {row['meta_key']: row['meta_value'] for row in metadata_rows}

        # Get sections for TOC and content
        cursor.execute("""
            SELECT id, section_key, parent_key, chapter_number, title, content_html,
                   display_order, toc_level, show_in_toc, word_count
            FROM thesis_sections
            WHERE is_active = 1
            ORDER BY display_order ASC
        """)
        sections = cursor.fetchall()

        # Build TOC
        toc = [
            {
                'section_key': s['section_key'],
                'chapter_number': s['chapter_number'] or '',
                'title': s['title'],
                'toc_level': s['toc_level']
            }
            for s in sections if s['show_in_toc']
        ]

        # Get references
        cursor.execute("""
            SELECT id, ref_key, authors, title, publication, year, doi, url, ref_type, formatted_citation
            FROM thesis_references
            WHERE is_active = 1
            ORDER BY display_order ASC
        """)
        references = cursor.fetchall()

        # Calculate stats
        total_words = sum(s['word_count'] or 0 for s in sections)

        cursor.close()
        conn.close()

        result = {
            'metadata': metadata,
            'toc': toc,
            'sections': sections,
            'figures': [],
            'tables': [],
            'references': references,
            'stats': {
                'total_sections': len(sections),
                'total_figures': 0,
                'total_tables': 0,
                'total_references': len(references),
                'total_words': total_words
            }
        }

        cache.set(cache_k, result, GUIDE_CACHE_TTL)

        return jsonify({'success': True, 'data': result, 'from_cache': False})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@dashboard_content_routes.route('/thesis/export/docx', methods=['GET'])
def export_thesis_docx():
    """Export thesis as DOCX file with professional academic formatting"""
    if not DOCX_AVAILABLE:
        return jsonify({'success': False, 'error': 'python-docx not installed'}), 500

    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        # Get metadata
        cursor.execute("SELECT meta_key, meta_value FROM thesis_metadata")
        metadata_rows = cursor.fetchall()
        metadata = {row['meta_key']: row['meta_value'] for row in metadata_rows}

        # Get sections with show_in_toc for TOC building
        cursor.execute("""
            SELECT section_key, chapter_number, title, content_html, toc_level, show_in_toc
            FROM thesis_sections
            WHERE is_active = 1
            ORDER BY display_order ASC
        """)
        sections = cursor.fetchall()

        # Get references
        cursor.execute("""
            SELECT authors, title, publication, year, doi, url, formatted_citation
            FROM thesis_references
            WHERE is_active = 1
            ORDER BY display_order ASC
        """)
        references = cursor.fetchall()

        cursor.close()
        conn.close()

        # Create document
        doc = Document()

        # Set up styles
        _setup_docx_styles(doc)

        # Set page margins (1 inch = 914400 EMUs, 1.25 inch for left binding)
        for section in doc.sections:
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # === COVER PAGE ===
        _add_cover_page(doc, metadata)

        # === TABLE OF CONTENTS ===
        doc.add_page_break()
        _add_table_of_contents(doc, sections)

        # === LIST OF FIGURES ===
        doc.add_page_break()
        _add_list_of_figures(doc)

        # === LIST OF TABLES ===
        doc.add_page_break()
        _add_list_of_tables(doc)

        # === MAIN CONTENT ===
        doc.add_page_break()
        for section in sections:
            _add_section_to_docx(doc, section)

        # === REFERENCES ===
        if references:
            doc.add_page_break()
            ref_heading = doc.add_heading('References', level=1)
            ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()
            for i, ref in enumerate(references, 1):
                citation = ref.get('formatted_citation') or _format_reference(ref)
                para = doc.add_paragraph()
                para.add_run(f'[{i}] ').bold = True
                para.add_run(citation)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5)
                para.paragraph_format.space_after = Pt(6)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"SSH_Guardian_Thesis_{metadata.get('student_id', 'Export')}.docx"

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def _setup_docx_styles(doc):
    """Set up document styles for professional thesis formatting"""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black for academic
        if i == 1:
            heading_style.font.size = Pt(16)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(10)
        elif i == 3:
            heading_style.font.size = Pt(13)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            heading_style.font.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(10)
            heading_style.paragraph_format.space_after = Pt(6)


def _add_cover_page(doc, metadata):
    """Add professional academic cover page"""
    # Add vertical spacing at top
    for _ in range(3):
        doc.add_paragraph()

    # Institution name
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(metadata.get('institution', 'ASIA PACIFIC UNIVERSITY OF TECHNOLOGY & INNOVATION'))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    # School/Faculty
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('School of Technology')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # Thesis Title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = metadata.get('title', 'SSH GUARDIAN: A MACHINE LEARNING-BASED SSH ANOMALY DETECTION SYSTEM FOR SMALL AND MEDIUM ENTERPRISES')
    run = para.add_run(title_text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # "A thesis submitted in partial fulfillment..."
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('A research project submitted in partial fulfillment of the requirements for the degree of')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()

    # Degree
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(metadata.get('degree', 'Master of Science in Cyber Security'))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # Student info
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Submitted by:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Student Name
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(metadata.get('student_name', 'MD SOHEL RANA'))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    # Student ID
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Student ID: {metadata.get('student_id', 'TP086217')}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    # Supervisor
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Supervised by:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(metadata.get('supervisor', 'Dr. Supervisor Name'))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    # Module Code
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Module: {metadata.get('module_code', 'CT095-6-M RMCE')}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Submission Date
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(metadata.get('submission_date', 'December 2025'))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def _add_table_of_contents(doc, sections):
    """Add Table of Contents with actual entries"""
    # TOC Title
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('TABLE OF CONTENTS')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Add TOC entries from sections
    page_num = 1  # Starting page estimate
    for section in sections:
        if not section.get('show_in_toc', True):
            continue

        toc_level = section.get('toc_level', 1)
        chapter_num = section.get('chapter_number', '')
        title = section.get('title', '')

        # Build entry text
        if chapter_num:
            if '.' in str(chapter_num):
                entry_text = f"{chapter_num} {title}"
            else:
                entry_text = f"Chapter {chapter_num}: {title}"
        else:
            entry_text = title

        # Create TOC entry paragraph
        para = doc.add_paragraph()

        # Indent based on level
        indent = (toc_level - 1) * 0.5
        para.paragraph_format.left_indent = Inches(indent)
        para.paragraph_format.space_after = Pt(4)

        # Add entry text
        run = para.add_run(entry_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Bold for chapter headings (level 1)
        if toc_level == 1:
            run.font.bold = True

    # Add note about updating
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run('[Note: Page numbers will be added after final formatting in Word]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def _add_list_of_figures(doc):
    """Add List of Figures placeholder"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('LIST OF FIGURES')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    # Figure entries (from thesis)
    figures = [
        ('Figure 1.1', 'Research Scope and Boundaries'),
        ('Figure 2.1', 'SSH Guardian Conceptual Framework'),
        ('Figure 2.2', 'Literature Review Taxonomy'),
        ('Figure 3.1', 'Design Science Research Methodology'),
        ('Figure 3.2', 'SSH Guardian Five-Layer Architecture'),
        ('Figure 3.3', 'Machine Learning Pipeline for SSH Attack Detection'),
        ('Figure 4.1', 'Dashboard Overview'),
        ('Figure 4.2', 'Confusion Matrix - Random Forest Model'),
        ('Figure 4.3', 'Feature Importance Analysis'),
    ]

    for fig_num, fig_title in figures:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f'{fig_num}: {fig_title}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def _add_list_of_tables(doc):
    """Add List of Tables placeholder"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('LIST OF TABLES')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    # Table entries (from thesis)
    tables = [
        ('Table 2.1', 'Comparative Analysis of SSH Security Solutions'),
        ('Table 2.2', 'Machine Learning Approaches for Intrusion Detection'),
        ('Table 3.1', 'Design Science Research Activities Mapping'),
        ('Table 3.2', 'Research Timeline'),
        ('Table 4.1', 'Agent Configuration Parameters'),
        ('Table 4.2', 'Complete Feature Set (50 Features by Category)'),
        ('Table 4.3', 'Random Forest Hyperparameters'),
        ('Table 4.4', 'Threat Intelligence API Summary'),
        ('Table 4.5', 'Cache TTL Configuration'),
        ('Table 4.6', 'Dataset Composition'),
        ('Table 4.7', 'Attack Simulation Scenarios'),
        ('Table 4.8', 'Model Comparison Results'),
        ('Table 4.9', 'Detection Rate by Attack Scenario'),
        ('Table 4.10', 'SSH Guardian vs Fail2ban Comparison'),
        ('Table 4.11', 'Resource Utilization Summary'),
    ]

    for tbl_num, tbl_title in tables:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f'{tbl_num}: {tbl_title}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def _add_section_to_docx(doc, section):
    """Add a thesis section to the document"""
    # Section heading
    chapter_num = section.get('chapter_number', '')
    title = section.get('title', 'Untitled')
    toc_level = section.get('toc_level', 1)

    # Build heading text
    if chapter_num:
        if '.' in str(chapter_num):
            # Sub-section (e.g., "1.1 Background")
            heading_text = f"{chapter_num} {title}"
        else:
            # Chapter heading (e.g., "Chapter 1: Introduction")
            heading_text = f"Chapter {chapter_num}: {title}"
    else:
        heading_text = title

    # Add heading with appropriate level (max 4 for Word)
    heading_level = min(toc_level, 4)
    heading = doc.add_heading(heading_text, level=heading_level)

    # Add page break before main chapter headings
    if toc_level == 1 and chapter_num and '.' not in str(chapter_num):
        # Insert page break before this heading
        run = heading.runs[0] if heading.runs else heading.add_run()
        run._element.addprevious(OxmlElement('w:br', {qn('w:type'): 'page'}))

    # Convert HTML content to Word paragraphs
    content_html = section.get('content_html', '')
    _html_to_docx(doc, content_html)


def _html_to_docx(doc, html_content):
    """Convert HTML content to Word document paragraphs with table and figure support"""
    if not html_content:
        return

    # Process content in segments - tables, figures, and text
    content = html_content

    # Patterns for tables and figures
    table_pattern = r'<table[^>]*>.*?</table>'
    figure_pattern = r'<figure[^>]*>.*?</figure>'

    # Also match table + following caption as a unit
    table_with_caption_pattern = r'(<table[^>]*>.*?</table>)\s*(<p[^>]*class=["\'][^"\']*table-caption[^"\']*["\'][^>]*>.*?</p>)?'

    # Find all tables (with potential captions) and figures
    all_matches = []

    # Find tables with their captions
    for match in re.finditer(table_with_caption_pattern, content, re.IGNORECASE | re.DOTALL):
        table_html = match.group(1)
        caption_html = match.group(2) if match.group(2) else ''
        combined = table_html + caption_html
        all_matches.append(('table', match.start(), match.end(), combined))

    # Find figures
    for match in re.finditer(figure_pattern, content, re.IGNORECASE | re.DOTALL):
        # Check if this figure overlaps with a table match
        overlaps = False
        for m in all_matches:
            if m[0] == 'table' and not (match.end() <= m[1] or match.start() >= m[2]):
                overlaps = True
                break
        if not overlaps:
            all_matches.append(('figure', match.start(), match.end(), match.group(0)))

    # Sort by position
    all_matches.sort(key=lambda x: x[1])

    # Process in order
    last_end = 0
    for match_type, start, end, match_content in all_matches:
        # Add text before this match
        if start > last_end:
            text_segment = content[last_end:start]
            _add_text_to_docx(doc, text_segment)

        # Add the table or figure
        if match_type == 'table':
            _add_html_table_to_docx(doc, match_content)
        elif match_type == 'figure':
            _add_figure_to_docx(doc, match_content)

        last_end = end

    # Add remaining text after last match
    if last_end < len(content):
        _add_text_to_docx(doc, content[last_end:])


def _add_text_to_docx(doc, html_text):
    """Add text content to Word document"""
    if not html_text or not html_text.strip():
        return

    # First, insert images where figure references are found
    text = _insert_figure_images(doc, html_text)


def _insert_figure_images(doc, html_text):
    """Insert figure images where figure references are found and return remaining text"""
    # Map figure numbers to image files and their captions
    # Note: Figure numbers in content map to available image files
    figure_data = {
        '1.1': ('figure_1_1_research_scope.png', 'Research Scope and Boundaries'),
        '2.1': ('figure_2_1_conceptual_framework.png', 'SSH Guardian Conceptual Framework'),
        '2.2': ('figure_2_2_literature_taxonomy.png', 'Literature Review Taxonomy'),
        '3.1': ('figure_3_1_dsr_methodology.png', 'Design Science Research Methodology'),
        '3.2': ('figure_3_2_system_architecture.png', 'SSH Guardian Five-Layer Architecture'),
        '3.3': ('figure_3_3_ml_pipeline.png', 'Machine Learning Pipeline for SSH Attack Detection'),
        # Figure 4.1 in content is Dashboard Overview (no image available)
        # Figure 4.2 in content is Feature Importance - mapped to feature_importance file
        '4.2': ('figure_4_3_feature_importance.png', 'Feature Importance Distribution (Top 15 Features)'),
        # Also include confusion matrix as Figure 4.3 in case it's referenced elsewhere
        '4.3': ('figure_4_2_confusion_matrix.png', 'Confusion Matrix - Random Forest Model'),
    }

    images_dir = PROJECT_ROOT / 'src' / 'dashboard' / 'static' / 'images' / 'thesis'

    # Pattern to find figure references - more flexible matching
    # Matches: "Figure X.Y:" or "Figure X.Y " at the start of a caption line
    figure_pattern = r'(Figure\s+(\d+\.\d+)\s*:?\s*)'

    # Track which figures have been inserted to avoid duplicates
    inserted_figures = set()

    # Find all figure references
    last_end = 0

    for match in re.finditer(figure_pattern, html_text, re.IGNORECASE):
        fig_num = match.group(2)

        # Only insert image if we haven't already and file exists
        if fig_num in figure_data and fig_num not in inserted_figures:
            img_filename, default_caption = figure_data[fig_num]
            img_path = images_dir / img_filename

            if img_path.exists():
                # Add text before this figure reference
                if match.start() > last_end:
                    text_before = html_text[last_end:match.start()]
                    _process_text_segment(doc, text_before)

                # Add the figure image
                try:
                    # Add centered paragraph for image
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))

                    # Add caption below image
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_text = f"Figure {fig_num}: {default_caption}"
                    run = caption_para.add_run(caption_text)
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.name = 'Times New Roman'

                    doc.add_paragraph()  # Space after figure

                    inserted_figures.add(fig_num)
                    last_end = match.end()
                    continue
                except Exception as e:
                    pass  # Fall through to add as text

        # If not inserting image, we'll process this as regular text
        # Don't update last_end - let it be processed as text

    # Process remaining text
    if last_end < len(html_text):
        _process_text_segment(doc, html_text[last_end:])

    return html_text  # Return for compatibility


def _process_text_segment(doc, html_text):
    """Process a segment of HTML text into Word paragraphs"""
    if not html_text or not html_text.strip():
        return

    text = html_text

    # Handle headings
    text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n### \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<h4[^>]*>(.*?)</h4>', r'\n#### \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<h5[^>]*>(.*?)</h5>', r'\n##### \1\n', text, flags=re.IGNORECASE | re.DOTALL)

    # Handle lists
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<[ou]l[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'</[ou]l>', '\n', text, flags=re.IGNORECASE)

    # Handle paragraphs
    text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)

    # Handle bold/italic
    text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', text, flags=re.IGNORECASE | re.DOTALL)

    # Handle div and span
    text = re.sub(r'<div[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<span[^>]*>(.*?)</span>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)

    # Remove remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)

    # Unescape HTML entities
    text = unescape(text)

    # Clean up whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    text = text.strip()

    if not text:
        return

    # Add paragraphs to document
    for para_text in text.split('\n\n'):
        para_text = para_text.strip()
        if not para_text:
            continue

        # Handle headings
        if para_text.startswith('### '):
            doc.add_heading(para_text[4:], level=3)
        elif para_text.startswith('#### '):
            doc.add_heading(para_text[5:], level=4)
        elif para_text.startswith('##### '):
            doc.add_heading(para_text[6:], level=4)
        else:
            # Handle bullet points as separate lines
            lines = para_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('• '):
                    # Bullet point
                    para = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(para, line[2:])
                else:
                    # Regular paragraph
                    para = doc.add_paragraph()
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    _add_formatted_text(para, line)


def _add_formatted_text(para, text):
    """Add text with bold/italic formatting to a paragraph"""
    # Handle bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic markers within
            subparts = re.split(r'(\*[^*]+\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                    run = para.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    para.add_run(subpart)


def _add_html_table_to_docx(doc, table_html, following_content=''):
    """Convert HTML table to Word table"""
    # Extract rows
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.IGNORECASE | re.DOTALL)
    if not rows:
        return

    # Parse all rows to determine dimensions
    parsed_rows = []
    max_cols = 0

    for row_html in rows:
        cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row_html, re.IGNORECASE | re.DOTALL)
        parsed_rows.append(cells)
        max_cols = max(max_cols, len(cells))

    if not parsed_rows or max_cols == 0:
        return

    # Create Word table
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fill table
    for row_idx, cells in enumerate(parsed_rows):
        row = table.rows[row_idx]
        for col_idx, cell_content in enumerate(cells):
            if col_idx < max_cols:
                cell = row.cells[col_idx]
                # Clean cell content
                cell_text = re.sub(r'<[^>]+>', '', cell_content)
                cell_text = unescape(cell_text).strip()
                cell.text = cell_text

                # Format header row
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Regular cells
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    # Add spacing after table
    doc.add_paragraph()

    # Look for table caption - check within table HTML and following content
    caption_text = None

    # Check for caption within table
    caption_match = re.search(r'<p[^>]*class=["\'][^"\']*table-caption[^"\']*["\'][^>]*>(.*?)</p>',
                              table_html, re.IGNORECASE | re.DOTALL)
    if caption_match:
        caption_text = caption_match.group(1)

    # Also check for traditional <caption> tag
    if not caption_text:
        caption_match = re.search(r'<caption[^>]*>(.*?)</caption>', table_html, re.IGNORECASE | re.DOTALL)
        if caption_match:
            caption_text = caption_match.group(1)

    if caption_text:
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clean_caption = re.sub(r'<[^>]+>', '', caption_text)
        clean_caption = unescape(clean_caption).strip()
        run = caption_para.add_run(clean_caption)
        run.font.size = Pt(10)
        run.font.italic = True

    doc.add_paragraph()


def _add_figure_to_docx(doc, figure_html):
    """Add figure with image and caption to document"""
    # Try to find image source
    img_match = re.search(r'<img[^>]*src=["\']([^"\']+)["\']', figure_html, re.IGNORECASE)
    caption_match = re.search(r'<figcaption[^>]*>(.*?)</figcaption>', figure_html, re.IGNORECASE | re.DOTALL)

    if img_match:
        img_src = img_match.group(1)

        # Try to load the image
        try:
            # Handle relative paths
            if img_src.startswith('/static/'):
                img_path = PROJECT_ROOT / 'src' / 'dashboard' / img_src.lstrip('/')
            elif img_src.startswith('../') or img_src.startswith('./'):
                img_path = PROJECT_ROOT / 'src' / 'dashboard' / 'static' / 'images' / 'thesis' / Path(img_src).name
            else:
                img_path = Path(img_src)

            if img_path.exists():
                # Add centered paragraph for image
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(img_path), width=Inches(5.5))
            else:
                # Image not found - add placeholder
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(f'[Image: {img_src}]')
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
        except Exception as e:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f'[Image could not be loaded: {img_src}]')
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

    # Add caption
    if caption_match:
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_text = re.sub(r'<[^>]+>', '', caption_match.group(1))
        caption_text = unescape(caption_text).strip()
        run = caption_para.add_run(caption_text)
        run.font.size = Pt(10)
        run.font.italic = True

    doc.add_paragraph()


def _format_reference(ref):
    """Format a reference in APA-like style"""
    parts = []
    if ref.get('authors'):
        parts.append(ref['authors'])
    if ref.get('year'):
        parts.append(f"({ref['year']})")
    if ref.get('title'):
        parts.append(ref['title'])
    if ref.get('publication'):
        parts.append(ref['publication'])
    if ref.get('doi'):
        parts.append(f"doi:{ref['doi']}")
    elif ref.get('url'):
        parts.append(ref['url'])
    return '. '.join(parts)


# ==============================================================================
# HELPER FUNCTIONS - Default Content
# ==============================================================================

def _get_default_step_content(step_key):
    """Get default HTML content for a guide step"""
    content = {
        'welcome': """
            <p>SSH Guardian is a <strong>lightweight SSH anomaly detection agent</strong> designed specifically for Small and Medium Enterprises (SMEs) operating on limited cloud infrastructure.</p>
            <p>This project is part of a Masters-level research thesis at <strong>Asia Pacific University of Technology & Innovation</strong>.</p>
            <h3>Key Features</h3>
            <ul>
                <li><strong>Real-time Monitoring</strong> - Track SSH authentication events as they happen</li>
                <li><strong>ML-Based Detection</strong> - Isolation Forest algorithm for anomaly detection</li>
                <li><strong>Threat Intelligence</strong> - Integration with AbuseIPDB and VirusTotal</li>
                <li><strong>Automated Response</strong> - Rule-based IP blocking via UFW</li>
                <li><strong>Resource Efficient</strong> - Minimal CPU/RAM footprint for SME servers</li>
            </ul>
        """,
        'architecture': """
            <p>SSH Guardian uses a <strong>multi-layered architecture</strong> to process and analyze authentication events efficiently:</p>
            <h3>Architecture Layers</h3>
            <ol>
                <li><strong>Data Collection Layer</strong> - Agent monitors SSH auth logs using inotify for real-time detection</li>
                <li><strong>Processing Layer</strong> - Events parsed, normalized, and enriched with GeoIP data</li>
                <li><strong>Analysis Layer</strong> - ML model + threat intelligence APIs evaluate risk score</li>
                <li><strong>Alert Layer</strong> - Notifications via Telegram, email, or webhooks</li>
                <li><strong>Storage Layer</strong> - MySQL database for events, statistics, and audit trail</li>
            </ol>
            <h3>Composite Risk Scoring</h3>
            <ul>
                <li><strong>35%</strong> - Threat Intelligence (AbuseIPDB, VirusTotal)</li>
                <li><strong>30%</strong> - ML Anomaly Score (Isolation Forest)</li>
                <li><strong>25%</strong> - Behavioral Patterns (velocity, uniqueness)</li>
                <li><strong>10%</strong> - Geographic Risk (high-risk regions)</li>
            </ul>
        """,
        'deployment': """
            <p>Deploy SSH Guardian agents on your cloud servers to start monitoring SSH authentication events.</p>
            <h3>Quick Installation</h3>
            <div style="background: var(--background); padding: 16px; border-radius: 8px; margin: 16px 0; overflow-x: auto; border: 1px solid var(--border);">
                <code style="color: var(--azure-blue); font-family: monospace;">curl -sSL https://ssh-guardian.rpu.solutions/install.sh | sudo bash</code>
            </div>
            <h3>Registration Process</h3>
            <ol>
                <li>Run the installer on your target server</li>
                <li>Agent generates a unique UUID and registers with this dashboard</li>
                <li>Navigate to <strong>Agents</strong> page and approve the new agent</li>
                <li>Agent begins sending real-time authentication events</li>
            </ol>
            <h3>System Requirements</h3>
            <ul>
                <li>Linux server (Ubuntu 20.04+, Debian 11+, CentOS 8+)</li>
                <li>Python 3.8 or higher</li>
                <li>Read access to /var/log/auth.log</li>
                <li>HTTPS access to dashboard API</li>
            </ul>
        """,
        'monitoring': """
            <p>The <strong>Live Events</strong> page shows all SSH authentication attempts across your monitored servers in real-time.</p>
            <h3>Event Types</h3>
            <ul>
                <li><span style="color: #ef4444; font-weight: bold;">Failed</span> - Invalid password or public key rejected</li>
                <li><span style="color: #10b981; font-weight: bold;">Successful</span> - Authenticated login session</li>
                <li><span style="color: #f59e0b; font-weight: bold;">Invalid User</span> - Username does not exist</li>
            </ul>
            <h3>Threat Analysis</h3>
            <p>Click <strong>"View Details"</strong> on any event to see:</p>
            <ul>
                <li>GeoIP location with map visualization</li>
                <li>AbuseIPDB reputation score and reports</li>
                <li>VirusTotal detection results</li>
                <li>ML anomaly assessment and confidence</li>
                <li>Historical patterns for this IP</li>
            </ul>
        """,
        'firewall': """
            <p>SSH Guardian integrates directly with <strong>UFW (Uncomplicated Firewall)</strong> for automated threat response.</p>
            <h3>Automatic Blocking Rules</h3>
            <p>Configure rules in <strong>Settings > Blocking Rules</strong> to automatically block IPs:</p>
            <ul>
                <li><strong>Failed Attempts</strong> - e.g., 5 failures in 10 minutes</li>
                <li><strong>ML Risk Score</strong> - Block when score exceeds threshold</li>
                <li><strong>AbuseIPDB Score</strong> - Block IPs with high abuse confidence</li>
                <li><strong>Geographic Restrictions</strong> - Block entire countries</li>
            </ul>
            <h3>Firewall Page Features</h3>
            <ul>
                <li>View all active blocks with expiration countdown</li>
                <li>Manually block/unblock individual IPs or CIDR ranges</li>
                <li>Configure UFW port rules (allow/deny services)</li>
                <li>Set block durations (15 min to permanent)</li>
            </ul>
        """,
        'intelligence': """
            <p>SSH Guardian combines multiple detection methods for comprehensive threat assessment.</p>
            <h3>Threat Intelligence APIs</h3>
            <ul>
                <li><strong>AbuseIPDB</strong> - Crowdsourced IP reputation with abuse confidence scores</li>
                <li><strong>VirusTotal</strong> - Multi-engine malware and malicious URL detection</li>
                <li><strong>FreeIPAPI</strong> - Geolocation, VPN detection, and proxy identification</li>
            </ul>
            <h3>Machine Learning Detection</h3>
            <p>The <strong>Isolation Forest</strong> algorithm detects anomalies based on:</p>
            <ul>
                <li><strong>Time Patterns</strong> - Unusual hours of activity</li>
                <li><strong>Geographic Anomalies</strong> - Login from new countries</li>
                <li><strong>Volume Patterns</strong> - Sudden spikes in attempts</li>
                <li><strong>Target Diversity</strong> - Multiple usernames from same IP</li>
            </ul>
        """
    }
    return content.get(step_key, '<p>Content not available for this step.</p>')


def _get_default_step_tips(step_key):
    """Get default tips for a guide step"""
    tips = {
        'welcome': 'Use the Quick Actions cards above to jump directly to key features, or continue through this guide to learn about each component.',
        'architecture': 'The architecture is designed for horizontal scalability - deploy multiple agents while centralizing monitoring in this dashboard.',
        'deployment': 'After installation, the agent appears in "Pending Approval" on the Agents page. Approve it to start receiving events.',
        'monitoring': 'Use filters to narrow events by date, type, agent, or risk level. High-risk events (score > 80) are highlighted.',
        'firewall': 'Enable Auto-Unblock with shorter durations (1-24 hours) to prevent permanent lockouts from legitimate users.',
        'intelligence': "The ML model learns your environment's baseline over time. Allow 1-2 weeks of data before relying heavily on ML scores."
    }
    return tips.get(step_key, '')


